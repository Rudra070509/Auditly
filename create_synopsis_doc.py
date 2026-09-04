import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_synopsis_docx(filename):
    doc = Document()

    # Page Setup: Standard Letter / A4 with 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles Configuration
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate-800

    # Colors
    NAVY = RGBColor(0x0F, 0x17, 0x2A)     # Deep Navy
    INDIGO = RGBColor(0x4F, 0x46, 0xE5)   # Primary Indigo
    DARK_SLATE = RGBColor(0x33, 0x41, 0x55)
    GRAY_TEXT = RGBColor(0x64, 0x74, 0x8B)

    # --- TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("PROJECT SYNOPSIS")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = INDIGO
    title_p.paragraph_format.space_after = Pt(2)

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run("AI-Powered Financial Audit Anomaly Detection Platform for Chartered Accountants & SME Audits")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = NAVY
    subtitle_p.paragraph_format.space_after = Pt(12)

    # Metadata Banner Table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        ("Domain:", "Financial Audit & Anomaly Detection"),
        ("Target Users:", "Chartered Accountants & SME Auditors"),
        ("Platform:", "Full-Stack Web App (React + Supabase)")
    ]

    for i, (label, val) in enumerate(meta_data):
        cell = meta_table.cell(0, i)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{label} ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK_SLATE
        r2 = p.add_run(val)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for section headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = INDIGO
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    # --- 1. EXECUTIVE SUMMARY ---
    add_heading_1("1. Executive Summary")
    p = doc.add_paragraph(
        "Small and Medium Enterprises (SMEs) represent a cornerstone of the economy, yet auditing SME accounting ledgers presents distinct operational hurdles for Chartered Accountants (CAs). Due to high transaction volumes and manual sample testing, irregular patterns such as backdated entries, circular round-tripping, GST tax claim mismatches, and year-end invoice padding frequently evade traditional statutory audit sampling."
    )
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph(
        "The AI-Powered Financial Audit Anomaly Detection Platform ('AuditPulse AI') addresses this vulnerability by delivering a automated, full-stack web application designed specifically for CA firms. By ingesting client general ledgers via standard Excel (.xlsx) and CSV files, the platform executes a 6-rule heuristic & statistical anomaly engine in real time. It calculates transparent Risk Scores (0–100), presents explainable AI natural-language reasoning, and automatically formats exportable Executive Audit Summaries for client management and working paper archives."
    )
    p2.paragraph_format.space_after = Pt(12)

    # --- 2. PROBLEM STATEMENT & OBJECTIVES ---
    add_heading_1("2. Problem Statement & Key Objectives")
    
    add_heading_2("2.1 Problem Statement")
    p = doc.add_paragraph(
        "Traditional statutory auditing relies heavily on selective manual sampling methods (e.g., inspecting 10-15% of vouchers). In SME environments characterized by high manual entry rates, this approach leaves critical blind spots:"
    )
    p.paragraph_format.space_after = Pt(4)

    bullet_points = [
        ("Undetected Round-Tripping:", " Circular fund transfers between sister concerns or related entities designed to artificially inflate turnover."),
        ("Backdated Accounting Vouchers:", " High-value entries created months after nominal invoice dates to absorb unallocated profits or expenses."),
        ("GST Credit Mismatches:", " Discrepancies between recorded ledger input tax credit and statutory GST rates (5%, 12%, 18%, 28%)."),
        ("Month-End Revenue/Expense Padding:", " Suspicious transaction volume spikes posted in the final 3 days of a financial month."),
        ("Statistical Outliers ($3\\sigma$):", " Unusual expenditure spikes exceeding 3 standard deviations above an account head's historical average.")
    ]

    for title, desc in bullet_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r_b = bp.add_run(title)
        r_b.font.bold = True
        r_b.font.color.rgb = NAVY
        r_d = bp.add_run(desc)
        r_d.font.color.rgb = DARK_SLATE

    add_heading_2("2.2 Key Project Objectives")
    objectives = [
        ("100% General Ledger Ingestion:", " Parse raw Excel/CSV trial balances and ledgers in client-side memory without manual data entry."),
        ("Real-Time 6-Rule Detection Engine:", " Automatically score every transaction voucher across 6 distinct financial risk heuristics."),
        ("Explainable AI Diagnosis:", " Synthesize clear, natural-language risk justifications for every flagged exception to assist CA engagement partners."),
        ("Auditor Working Paper Workflow:", " Provide interactive audit review actions (Approve, Flag for Follow-up, Dismiss) with custom auditor comments."),
        ("Executive Summary Generation:", " Render print-ready, professional PDF executive summaries bearing formal CA firm letterheads.")
    ]

    for title, desc in objectives:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r_b = bp.add_run(title)
        r_b.font.bold = True
        r_b.font.color.rgb = INDIGO
        r_d = bp.add_run(desc)

    # --- 3. SYSTEM ARCHITECTURE & TECH STACK ---
    add_heading_1("3. System Architecture & Technical Stack")
    p = doc.add_paragraph("The platform is architected as a modern, decoupled full-stack single-page application (SPA):")
    p.paragraph_format.space_after = Pt(8)

    # Tech Stack Table
    stack_table = doc.add_table(rows=6, cols=3)
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stack_table.autofit = False

    headers = ["Layer", "Technology Selection", "Role & Rationale"]
    hdr_row = stack_table.rows[0]
    for j, text in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    stack_rows = [
        ("Frontend UI", "React (Vite) + Tailwind CSS", "Responsive dark/light dashboard UI with instant state management"),
        ("Data Visualization", "Recharts", "Interactive trend line charts, category bar charts & exposure cards"),
        ("Excel Parsing", "SheetJS (XLSX) & PapaParse", "Client-side binary parsing of .xlsx, .xls, and .csv ledger exports"),
        ("Database & Storage", "Supabase (PostgreSQL)", "Cloud persistence for audit_transactions table with LocalStorage fallback"),
        ("Anomaly Engine", "TypeScript / JS Pure Functions", "Client-side mathematical engine evaluating 6 risk rules & 3-Sigma math")
    ]

    for i, row in enumerate(stack_rows, start=1):
        tr = stack_table.rows[i]
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        for j, text in enumerate(row):
            cell = tr.cells[j]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if j == 0:
                r.font.bold = True
                r.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 4. CORE ANOMALY DETECTION ENGINE (6 HEURISTICS) ---
    add_heading_1("4. Core Anomaly Detection Heuristics")
    p = doc.add_paragraph(
        "The anomaly engine processes retrieved transactions through 6 mathematical & rule-based algorithms:"
    )
    p.paragraph_format.space_after = Pt(8)

    # Heuristics Table
    rule_table = doc.add_table(rows=7, cols=4)
    rule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rule_table.autofit = False

    r_headers = ["Rule ID", "Anomaly Type", "Detection Logic", "Score Impact"]
    r_hdr_row = rule_table.rows[0]
    for j, text in enumerate(r_headers):
        cell = r_hdr_row.cells[j]
        set_cell_background(cell, "4F46E5")
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rules_data = [
        ("Rule A", "Duplicate Transactions", "Identical amount + same account head within a 48-hour date window", "+35 (High)"),
        ("Rule B", "Backdated Entries", "System posting date is >30 days later than nominal transaction date", "+25 to +40"),
        ("Rule C", "Round-Tripping / Circular", "Matching debit/credit fund flows between related entities within 5 days", "+45 (High)"),
        ("Rule D", "GST-to-Book Mismatch", "Recorded ledger GST diverges >5% from statutory rates (5%, 12%, 18%, 28%)", "+20 to +30"),
        ("Rule E", "Month-End Spikes", "Voucher value >2.5x mean posted in final 3 days of financial month", "+25 (Medium)"),
        ("Rule F", "3-Sigma Statistical Outlier", "Transaction value exceeds 3 standard deviations (Z-Score >= 3.0σ) above mean", "+40 (High)")
    ]

    for i, row in enumerate(rules_data, start=1):
        tr = rule_table.rows[i]
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        for j, text in enumerate(row):
            cell = tr.cells[j]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            if j == 0 or j == 1:
                r.font.bold = True
                r.font.color.rgb = NAVY
            if j == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) if "High" in text else INDIGO

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 5. APPLICATION MODULES & USER WORKFLOW ---
    add_heading_1("5. Application Modules & User Workflow")

    modules = [
        ("5.1 Data Ingestion & Column Header Mapping Screen",
         "Features a drag-and-drop file upload zone accepting .xlsx and .csv files. Triggers an interactive Column Mapping Confirmation Modal allowing auditors to align custom client Excel headers (e.g., 'Vouch No', 'Trx Date', 'Particulars') to canonical database fields with real-time sample data previews."),

        ("5.2 Risk Overview Analytics Dashboard",
         "Displays 4 executive KPI cards (Total Vouchers Analyzed, Total Anomalies Detected, High-Risk Exposure ₹, Pending Reviews). Integrates Recharts visualizers including a Timeline Line Chart (Volume vs Flagged Anomalies) and a Risk Distribution Bar Chart across all 6 categories."),

        ("5.3 Anomaly Explorer & Ledger Working Paper Table",
         "Provides a filterable table with real-time search, risk tier badges (High, Medium, Low, Pass), and anomaly tags. Clicking any entry opens a Detail Drawer Modal showing complete voucher metadata, rule breakdown, explainable AI diagnosis, and auditor action buttons (Approve, Flag for Working Paper, Dismiss with Notes)."),

        ("5.4 Audit Report Generator & Executive Summary",
         "Generates a printable, formal Executive Audit Summary featuring CA firm letterhead ('R. MEHTA & CO. Chartered Accountants'), auditor executive opinion statements, high-risk exception tables, digital stamp area, and partner signature block formatted for PDF export.")
    ]

    for title, desc in modules:
        add_heading_2(title)
        p = doc.add_paragraph(desc)
        p.paragraph_format.space_after = Pt(6)

    # --- 6. BUSINESS IMPACT & CONCLUSION ---
    add_heading_1("6. Business Impact & Conclusion")
    p = doc.add_paragraph(
        "The AI-Powered Financial Audit Anomaly Detection Platform transforms SME statutory audits from a slow, sample-based manual process into a continuous, data-driven workflow. Key quantifiable benefits include:"
    )
    p.paragraph_format.space_after = Pt(6)

    benefits = [
        ("Time Efficiency:", " Reduces preliminary voucher scanning time from days to seconds."),
        ("Enhanced Audit Quality:", " Achieves 100% ledger coverage rather than relying on 10-15% manual sample testing."),
        ("Risk Mitigation:", " Early identification of round-tripping and GST mismatches prevents regulatory penalties."),
        ("Professional Deliverables:", " Generates clean, executive-ready PDF audit summaries for SME management boards.")
    ]

    for title, desc in benefits:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r_b = bp.add_run(title)
        r_b.font.bold = True
        r_b.font.color.rgb = INDIGO
        r_d = bp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Sign-off Block
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s1 = sign_p.add_run("Prepared & Submitted by:\n")
    r_s1.font.size = Pt(9.5)
    r_s1.font.color.rgb = GRAY_TEXT
    r_s2 = sign_p.add_run("AuditPulse AI Engineering Team\n")
    r_s2.font.bold = True
    r_s2.font.size = Pt(11)
    r_s2.font.color.rgb = NAVY
    r_s3 = sign_p.add_run("Chartered Accountant Technology Solutions Division")
    r_s3.font.size = Pt(9.5)
    r_s3.font.color.rgb = GRAY_TEXT

    doc.save(filename)
    print(f"Synopsis successfully saved to {filename}")

if __name__ == "__main__":
    create_synopsis_docx("/Users/Admin/Desktop/Hackathon/AI_Powered_Financial_Audit_Anomaly_Detection_Platform_Synopsis.docx")
